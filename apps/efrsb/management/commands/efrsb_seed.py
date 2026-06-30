"""Идемпотентный сидинг ЕФРСБ.

  • базовая .docx-заготовка сообщения (kind=efrsb) → S3 + DocumentTemplate;
  • ПОЛНЫЙ каталог типов сообщений (EfrsbMessageType) из официального справочника
    ЕФРСБ (снапшот apps/efrsb/reference_data/message_types.json, метод read-API
    v1/reference-books/message-types) — код = официальный код, api_type = код;
  • типы отчётов (Приложение 3 спецификации) как api_kind=report;
  • типы лога (EventType/ActionType) для событийки.

🛑 НЕ входит в deploy-handler — гонять вручную (как procedure_seed):
    docker exec siricrm-web-1 python manage.py efrsb_seed

🛑 Каталог грузится в статусе DRAFT (is_draft=True). Поля, специфичные для БФЛ
   (applicable_kinds, sets_efrsb_date, привязка пер-типового шаблона, сроки) —
   подтверждаются с АУ и правятся в Справочниках; повторный сид их НЕ затирает.

Обновление снапшота (на dev, с demo-контура), если ЕФРСБ добавит типы:
    python manage.py efrsb_pull_reference   # см. отдельную команду
"""
from __future__ import annotations

import io
import json
import os

from django.core.management.base import BaseCommand

from apps.afd.models import DocumentTemplate
from apps.efrsb.models import EfrsbMessageType
from apps.files.models import StoredFile
from apps.files.s3_utils import upload_file_to_s3

_DOCX_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_TEMPLATE_NAME = "Сообщение ЕФРСБ — базовая заготовка (черновик)"

_REF_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "reference_data")
_MESSAGE_TYPES_JSON = os.path.join(_REF_DIR, "message_types.json")

# Нейтральная структурная «рыба» сообщения. Содержательную часть («Текст сообщения»)
# вводит АУ при формировании; реквизиты должника/АУ/дела подставляются из CRM.
# Это ЗАГОТОВКА — АУ заводит полноценные пер-типовые шаблоны в разделе АФД.
_SKELETON_PARAS = [
    ("{Заголовок}", "center", True),
    ("", "both", False),
    ("Финансовый управляющий {ФИО Финансовый управляющий} (ИНН {ИНН АУ}, СНИЛС {СНИЛС АУ}, "
     "адрес для корреспонденции: {Адрес арбитражного управляющего}), член {Реквизиты СРО}, "
     "действующий в деле о банкротстве № {номер дела} ({арбитражный суд}) в отношении "
     "{Фамилия} {Имя} {Отчество} (дата рождения {дата рождения}, место рождения {место рождения}, "
     "ИНН {ИНН}, СНИЛС {СНИЛС}, адрес регистрации: {адрес регистрации}), сообщает:", "both", False),
    ("", "both", False),
    ("{Текст сообщения}", "both", False),
    ("", "both", False),
    ("Дата: {дата}", "both", False),
    ("Финансовый управляющий ________________ {ФамилияИО АУ}", "both", False),
]

# Типы отчётов ФУ — Приложение 3 спецификации (нет в reference-books/message-types).
# (code, name, is_old)
_REPORT_TYPES = [
    ("Final", "Финальный отчёт", False),
    ("Final2", "Финальный отчёт", False),
    ("Annulment", "Аннулирование ранее опубликованного отчёта", False),
    ("Annulment2", "Аннулирование ранее опубликованного отчёта", False),
    ("SignificantEvent", "Отчёт по существенным фактам (устаревший)", True),
    ("Periodic", "Периодический отчёт (устаревший)", True),
]

# 🛑 «Вводные» типы — проставляют Procedure.publication_efrsb_date. ЗАГОТОВКА, TO CONFIRM:
# судебный акт о введении процедуры публикуется в ЕФРСБ как «Сообщение о судебном акте».
_SETS_DATE_CODES = {"ArbitralDecree"}

# 🛑 БФЛ-шортлист (релевантно банкротству ГРАЖДАН) — первичная разметка, TO CONFIRM с АУ.
# По умолчанию в селекторе показываются только эти; остальные активные — по «Показать все».
# Исключены типы для юрлиц/банков/застройщиков/субсидиарки/внесудебного банкротства.
_BFL_CODES = {
    # старт дела / акты / реестр требований
    "ArbitralDecree", "CourtAcceptanceStatement", "DemandAnnouncement",
    "CreditorsDemandRegistered", "FinancialStateInformation",
    "DeliberateBankruptcy", "CancelDeliberateBankruptcy", "ChangeDeliberateBankruptcy",
    # собрания кредиторов
    "Meeting2", "MeetingResult",
    # имущество: опись/оценка/реализация
    "PropertyInventoryResult", "ProcessInventoryDebtor", "AssessmentReport",
    "CourseOfSalePersonProperty", "SaleOrderPledgedProperty2", "MortgageSaleExclusion",
    "TransferAssertsForImplementation",
    # торги
    "Auction2", "ChangeAuction2", "CancelAuctionTradeResult", "TradeResult",
    "SaleContractResult2",
    # расчёты с кредиторами
    "OrderAndTimingCalculations", "StartSettlement", "ProcedureGrantingIndemnity",
    # реструктуризация
    "ViewDraftRestructuringPlan", "ViewExecRestructuringPlan",
    # оспаривание сделок гражданина
    "DealInvalid2", "DealInvalidResult2",
    # служебные
    "Annul", "Other",
    # отчёты ФУ
    "Final", "Final2", "Annulment", "Annulment2",
}

# Старые наши «ручные» коды из первой версии сида (до перехода на официальные коды) —
# удаляем при наличии, чтобы не было дублей концепций.
_LEGACY_CODES = [
    "arbitral_decree", "meeting_creditors", "meeting_result", "inventory",
    "evaluation", "auction", "trade_result", "sale_contract", "demand_announcement",
    "fin_state", "deliberate_bankruptcy", "calc_order", "restr_plan_view",
    "final_report", "other",
]

_EVENT_TYPES = [
    ("efrsb_published_detected", "Публикация в ЕФРСБ обнаружена", "system", False, ""),
    ("efrsb_violation", "Публикация ЕФРСБ с нарушением срока", "system", True,
     "ЕФРСБ отметил публикацию как сделанную с нарушением срока — проверьте."),
]
_ACTION_TYPES = [
    ("efrsb_text_generated", "Сформирован текст сообщения ЕФРСБ", False),
]


def _build_skeleton_docx() -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    align_map = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    doc = Document()
    for text, align, bold in _SKELETON_PARAS:
        p = doc.add_paragraph()
        p.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
        run = p.add_run(text)
        run.bold = bold
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


class Command(BaseCommand):
    help = "Идемпотентный сидинг ЕФРСБ (заготовка-шаблон + ПОЛНЫЙ каталог типов из справочника + типы лога)."

    def handle(self, *args, **opts):
        self._cleanup_legacy()
        tpl = self._seed_template()
        self._load_message_types(tpl)
        self._load_report_types(tpl)
        self._seed_log_types()
        self.stdout.write(self.style.SUCCESS("ЕФРСБ: сидинг завершён."))

    def _cleanup_legacy(self):
        qs = EfrsbMessageType.objects.filter(code__in=_LEGACY_CODES)
        n = qs.count()
        if n:
            qs.delete()
            self.stdout.write(self.style.WARNING(
                f"• Удалены {n} старых «ручных» типов (перешли на официальные коды ЕФРСБ)."))

    def _seed_template(self) -> DocumentTemplate | None:
        tpl = DocumentTemplate.objects.filter(
            kind=DocumentTemplate.KIND_EFRSB, name=_TEMPLATE_NAME
        ).first()
        if tpl:
            self.stdout.write("• Заготовка-шаблон ЕФРСБ уже есть — пропуск.")
            return tpl
        try:
            data = _build_skeleton_docx()
        except Exception as exc:  # noqa: BLE001
            self.stdout.write(self.style.ERROR(
                f"• Не удалось собрать заготовку .docx ({exc}). "
                "Заведите шаблон вручную в разделе АФД."))
            return None
        bucket, key = upload_file_to_s3(
            data, prefix="afd/efrsb_templates",
            filename="efrsb_message_skeleton.docx", content_type=_DOCX_CT)
        sf = StoredFile.objects.create(
            bucket=bucket, key=key, filename=f"{_TEMPLATE_NAME}.docx",
            content_type=_DOCX_CT, size=len(data))
        tpl = DocumentTemplate.objects.create(
            name=_TEMPLATE_NAME,
            kind=DocumentTemplate.KIND_EFRSB,
            stored_file=sf,
            description=(
                "Черновик-заготовка сообщения ЕФРСБ. Плейсхолдеры: {Заголовок} "
                "{Текст сообщения} {Фамилия} {Имя} {Отчество} {дата рождения} "
                "{место рождения} {ИНН} {СНИЛС} {адрес регистрации} "
                "{ФИО Финансовый управляющий} {ФамилияИО АУ} {ИНН АУ} {СНИЛС АУ} "
                "{Адрес арбитражного управляющего} {Реквизиты СРО} {арбитражный суд} "
                "{номер дела} {дата}. Замените на пер-типовые шаблоны в разделе АФД."))
        self.stdout.write(self.style.WARNING(
            "• Создана базовая заготовка-шаблон ЕФРСБ (черновик)."))
        return tpl

    def _upsert(self, *, code, name, api_kind, is_old, tpl):
        """UPSERT типа по коду. На создании — дефолты + шаблон; на обновлении —
        только name/is_active (официальная истина), БФЛ-поля не трогаем."""
        obj = EfrsbMessageType.objects.filter(code=code).first()
        if obj is None:
            EfrsbMessageType.objects.create(
                code=code, name=name, api_kind=api_kind, api_type=code,
                is_active=(not is_old), is_draft=True,
                is_bfl=(code in _BFL_CODES),
                sets_efrsb_date=(code in _SETS_DATE_CODES),
                template=tpl if tpl else None,
            )
            return "created"
        # обновляем «истину» из справочника, не затирая правки АУ
        changed = False
        if obj.name != name:
            obj.name, changed = name, True
        if obj.is_active != (not is_old):
            obj.is_active, changed = (not is_old), True
        if obj.api_type != code:
            obj.api_type, changed = code, True
        if obj.template_id is None and tpl:
            obj.template, changed = tpl, True
        if changed:
            obj.save()
            return "updated"
        return "kept"

    def _load_message_types(self, tpl):
        with open(_MESSAGE_TYPES_JSON, encoding="utf-8") as f:
            items = json.load(f)
        stats = {"created": 0, "updated": 0, "kept": 0}
        for it in items:
            res = self._upsert(code=it["code"], name=it["name"],
                               api_kind=EfrsbMessageType.API_KIND_MESSAGE,
                               is_old=bool(it.get("isOld")), tpl=tpl)
            stats[res] += 1
        self.stdout.write(self.style.SUCCESS(
            f"• Типы сообщений ({len(items)}): создано {stats['created']}, "
            f"обновлено {stats['updated']}, без изменений {stats['kept']} (DRAFT)."))

    def _load_report_types(self, tpl):
        stats = {"created": 0, "updated": 0, "kept": 0}
        for code, name, is_old in _REPORT_TYPES:
            res = self._upsert(code=code, name=name,
                               api_kind=EfrsbMessageType.API_KIND_REPORT,
                               is_old=is_old, tpl=tpl)
            stats[res] += 1
        self.stdout.write(self.style.SUCCESS(
            f"• Типы отчётов ({len(_REPORT_TYPES)}): создано {stats['created']}, "
            f"обновлено {stats['updated']}, без изменений {stats['kept']}."))

    def _seed_log_types(self):
        from apps.crm.models import ActionType, EventType
        for code, name, source, notifies, hint in _EVENT_TYPES:
            EventType.objects.update_or_create(
                code=code,
                defaults={"name": name, "source": source, "is_system": True,
                          "is_manual": False, "is_active": True,
                          "notifies": notifies, "notify_hint": hint})
        for code, name, notifies in _ACTION_TYPES:
            ActionType.objects.update_or_create(
                code=code,
                defaults={"name": name, "is_system": True, "is_manual": False,
                          "is_active": True, "notifies": notifies})
        self.stdout.write("• Типы лога (EventType/ActionType) ЕФРСБ обновлены.")
