"""Генератор ТЕКСТА сообщения ЕФРСБ.

Поток (как в ЛК fedresurs): выбрали тип/подтип → проверили недостающие данные →
«Сгенерировать текст» подставляет CRM-данные в текстовый шаблон типа/подтипа →
текст правится вручную → «Сохранить»: текст ложится в публикацию, из него же
собираются .docx/.pdf и подшиваются в папку клиента «Публикации ЕФРСБ».

Шаблон текста живёт в `EfrsbMessageSubtype.text_template` (приоритет) или
`EfrsbMessageType.text_template`. Плейсхолдеры — {Русские ключи} (см. build_context).
Текст нужен для РУЧНОЙ публикации АУ в ЛК fedresurs (авто-публикации у оператора нет).
"""
from __future__ import annotations

import io
import logging
import re

from django.utils import timezone

from apps.afd.pdf_utils import docx_to_pdf
from apps.crm import client_log
from apps.files.folder_utils import _mk, get_or_create_root
from apps.files.models import ClientFile, StoredFile
from apps.files.s3_utils import upload_file_to_s3
from apps.procedure.request_documents import _debtor_address, _fmt, _spouse_data

log = logging.getLogger(__name__)
DOCX_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

_PH_RE = re.compile(r"\{[^{}\n]+\}")


class EfrsbGenError(RuntimeError):
    pass


# ── контекст ────────────────────────────────────────────────────────────────

def _procedure_for(case, procedure=None):
    """Процедура для реквизитов: явная → текущая → последняя с назначенным АУ."""
    if procedure is not None:
        return procedure
    if case.current_procedure_id and case.current_procedure:
        return case.current_procedure
    return (case.procedures.exclude(arbitr_manager=None).order_by("-order").first()
            or case.procedures.order_by("-order").first())


_CASE_PREFIX_RE = re.compile(r"^\s*[АA](\d{1,2})\s*-")


def _court_region(service, arb):
    """Суд определяем ПО НОМЕРУ ДЕЛА: префикс «А12-…» → Region.arbitr_code «А12»
    → Арбитражный суд Волгоградской области. Фолбэк — регион услуги.

    🛑 Нельзя брать `ArbitrCase.court_name` напрямую: kad-парсер пишет туда
    ТЕКУЩУЮ инстанцию вместе с судьёй («13 арбитражный апелляционный суд /
    Фуркало О. В.»), а в сообщении нужен суд первой инстанции.
    """
    from apps.crm.models import Region
    number = (arb.case_number or "").strip() if arb is not None else ""
    m = _CASE_PREFIX_RE.match(number)
    if m:
        code = "А%02d" % int(m.group(1))  # А12, А07 — кириллическая «А»
        region = Region.objects.filter(arbitr_code__iexact=code).order_by("number").first()
        if region is not None:
            return region
    return getattr(service, "region", None)


def _court_genitive(name: str) -> str:
    """«Арбитражный суд Волгоградской области» → «Арбитражного суда Волгоградской
    области» (для оборота «в помещении …»). Все АС субъектов названы по одному
    шаблону, поэтому склонение сводится к замене первых двух слов."""
    n = (name or "").strip()
    if n.startswith("Арбитражный суд "):
        return "Арбитражного суда " + n[len("Арбитражный суд "):]
    return n


def _court(service, arb) -> tuple[str, str]:
    """(наименование суда, адрес суда)."""
    region = _court_region(service, arb)
    name = (region.court_name or "").strip() if region is not None else ""
    if not name and arb is not None:
        # Фолбэк: из kad, но без «/ Судья И. О.» в хвосте.
        name = (arb.court_name or "").split("/")[0].strip()
    address = ""
    if region is not None and getattr(region, "court_address_id", None):
        address = (region.court_address.result or "").strip()
    return name, address


def _sro_full(am) -> str:
    """«Ассоциации … "СОЛИДАРНОСТЬ" (ИНН: …, ОГРН: …, адрес: …)».

    🛑 Приоритет — ручной `ArbitrationManager.sro_text`: из ЕГРЮЛ название приходит
    КАПСОМ и в именительном падеже («АССОЦИАЦИЯ …»), а в сообщении нужна нормальная
    формулировка в родительном («член Ассоциации …»). Автосклонять произвольные
    названия нельзя — АУ один раз вписывает готовый текст в реквизиты СРО.
    """
    if am is None:
        return ""
    if (am.sro_text or "").strip():
        return am.sro_text.strip()
    sro = am.sro if am.sro_id else None
    if sro is None:
        return ""
    meta = []
    if sro.inn:
        meta.append(f"ИНН: {sro.inn}")
    if sro.ogrn:
        meta.append(f"ОГРН: {sro.ogrn}")
    addr = (sro.legal_address or "").strip()
    if addr:
        meta.append(f"адрес: {addr}")
    name = (sro.name or sro.short_name or "").strip()
    return f"{name} ({', '.join(meta)})" if meta else name


def _fu_approved(am) -> str:
    """«утверждён» / «утверждена» — согласование по полу управляющего.

    Пол берём по отчеству (женские -вна/-чна, иначе мужская форма). В справочнике
    АУ пола нет, а заводить ради одного слова избыточно. 🛑 Эвристика: если
    промахнулась — АУ правит формулировку в тексте перед сохранением.
    (Та же логика, что в apps.kommersant.generator — держим согласованно.)
    """
    patronymic = (getattr(am, "patronymic", "") or "").strip().lower()
    return "утверждена" if patronymic.endswith(("вна", "чна")) else "утверждён"


def build_context(case, *, message_type=None, subtype=None, procedure=None, overrides=None) -> dict:
    """Плоский dict плейсхолдеров. Публикация не требуется — можно звать до её создания."""
    overrides = dict(overrides or {})
    client = case.service.client
    proc = _procedure_for(case, procedure)
    am = proc.arbitr_manager if (proc and proc.arbitr_manager_id) else None
    arb = getattr(case.service, "arbitr_case", None)
    idx, addr_reg = _debtor_address(client)
    fio = " ".join(filter(None, [client.last_name, client.first_name, client.patronymic])).strip()
    court_name, court_address = _court(case.service, arb)

    ctx = {
        # Тип/подтип
        "Тип сообщения": (message_type.name if message_type else ""),
        "Тип судебного акта": (subtype.name if subtype else ""),
        # Должник
        "ФИО должника": fio,
        "Фамилия": client.last_name or "", "Имя": client.first_name or "",
        "Отчество": client.patronymic or "",
        "дата рождения": _fmt(client.birth_date),
        "место рождения": client.birth_place or "",
        "ИНН": client.inn or "", "СНИЛС": client.snils or "",
        "индекс": idx, "адрес регистрации": addr_reg,
        # Финуправляющий (публикатор)
        "ФИО Финансовый управляющий": am.full_fio if am else "",
        "утверждён ФУ": _fu_approved(am),   # утверждён / утверждена — по полу
        "ФамилияИО АУ": am.short_fio if am else "",
        "ИНН АУ": am.inn if am else "", "СНИЛС АУ": am.snils if am else "",
        "Адрес арбитражного управляющего": am.corr_address if am else "",
        "Телефон арбитражного": am.phone if am else "",
        "email арбитражного": am.email if am else "",
        "Реквизиты СРО": am.sro_display if am else "",
        "СРО полностью": _sro_full(am),
        # Дело / суд / процедура (суд — по номеру дела, см. _court)
        "арбитражный суд": court_name,
        "арбитражного суда": _court_genitive(court_name),  # «в помещении …»
        "адрес суда": court_address,
        "номер дела": (arb.case_number if arb else ""),
        "вид процедуры": (proc.get_kind_display() if proc else ""),
        "дата решения": _fmt(proc.intro_date) if proc else "",
        "дата завершения": _fmt(proc.end_date) if proc else "",  # определение об окончании
        "срок процедуры": (str(proc.term_months) if proc and proc.term_months else ""),
        "дата следующего заседания": _fmt(proc.next_hearing_date) if proc else "",
        "дата публикации ЕФРСБ": _fmt(proc.publication_efrsb_date) if proc else "",
        # Прочее
        "данные на супруга": _spouse_data(client.spouse),
        "дата": _fmt(timezone.localdate()),
    }
    ctx.update({k: v for k, v in overrides.items() if v is not None})
    return ctx


# ── шаблон текста ───────────────────────────────────────────────────────────

def resolve_text_template(message_type, subtype) -> str:
    """Шаблон подтипа приоритетнее шаблона типа."""
    if subtype is not None and (subtype.text_template or "").strip():
        return subtype.text_template
    if message_type is not None and (message_type.text_template or "").strip():
        return message_type.text_template
    return ""


def placeholders_in(template: str) -> list[str]:
    out, seen = [], set()
    for m in _PH_RE.finditer(template or ""):
        key = m.group(0)[1:-1]
        if key not in seen:
            seen.add(key)
            out.append(key)
    return out


def substitute(template: str, ctx: dict) -> str:
    def _s(m):
        key = m.group(0)[1:-1]
        if key in ctx:
            v = ctx[key]
            return "" if v is None else str(v)
        return m.group(0)  # неизвестный плейсхолдер оставляем видимым
    return _PH_RE.sub(_s, template or "")


def render_message_text(case, message_type, subtype, *, procedure=None, overrides=None) -> str:
    """Сгенерировать текст сообщения по шаблону типа/подтипа."""
    tpl = resolve_text_template(message_type, subtype)
    if not tpl.strip():
        raise EfrsbGenError(
            "У типа/подтипа не задан шаблон текста. Заполните «Шаблон текста сообщения» "
            "в справочнике (или введите текст вручную)."
        )
    ctx = build_context(case, message_type=message_type, subtype=subtype,
                        procedure=procedure, overrides=overrides)
    return substitute(tpl, ctx).strip()


# ── проверка данных (показываем ТОЛЬКО проблемы) ────────────────────────────

_LABELS = {
    "ФИО должника": "ФИО должника",
    "дата рождения": "Дата рождения должника",
    "место рождения": "Место рождения должника",
    "ИНН": "ИНН должника", "СНИЛС": "СНИЛС должника",
    "адрес регистрации": "Адрес регистрации должника",
    "ФИО Финансовый управляющий": "ФИО финуправляющего",
    "ИНН АУ": "ИНН финуправляющего", "СНИЛС АУ": "СНИЛС финуправляющего",
    "Адрес арбитражного управляющего": "Адрес корреспонденции ФУ",
    "email арбитражного": "E-mail финуправляющего",
    "Реквизиты СРО": "СРО", "СРО полностью": "СРО (наименование, ИНН, ОГРН, адрес)",
    "арбитражный суд": "Арбитражный суд", "арбитражного суда": "Арбитражный суд",
    "адрес суда": "Адрес арбитражного суда",
    "номер дела": "Номер дела",
    "дата решения": "Дата решения суда (введение процедуры)",
    "дата завершения": "Дата определения о завершении процедуры",
    "срок процедуры": "Срок процедуры, мес.",
    "дата следующего заседания": "Дата следующего судебного заседания",
}
_DIGIT_RULES = {"ИНН": (10, 12), "ИНН АУ": (10, 12), "СНИЛС": (11,), "СНИЛС АУ": (11,)}


def check_problems(case, message_type, subtype, *, procedure=None) -> list[dict]:
    """Только НЕзаполненные/невалидные данные, нужные шаблону. [] — всё в порядке."""
    tpl = resolve_text_template(message_type, subtype)
    if not tpl.strip():
        return [{"label": "Шаблон текста сообщения",
                 "note": "не задан у типа/подтипа — текст можно ввести вручную"}]
    ctx = build_context(case, message_type=message_type, subtype=subtype, procedure=procedure)
    problems = []
    for key in placeholders_in(tpl):
        val = str(ctx.get(key) or "").strip()
        label = _LABELS.get(key, key)
        if not val:
            problems.append({"label": label, "note": "не заполнено"})
            continue
        rule = _DIGIT_RULES.get(key)
        if rule and len(re.sub(r"\D", "", val)) not in rule:
            problems.append({"label": label, "note": "неверный формат"})
            continue
        if key == "email арбитражного" and "@" not in val:
            problems.append({"label": label, "note": "неверный e-mail"})
    return problems


# ── сохранение: текст → публикация + .docx/.pdf в папку клиента ─────────────

def _text_to_docx(title: str, text: str) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.add_run(title).bold = True
    doc.add_paragraph("")
    for line in (text or "").split("\n"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(line)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _store(file_bytes, *, filename, content_type):
    bucket, key = upload_file_to_s3(
        file_bytes, prefix="efrsb/messages", filename=filename, content_type=content_type,
    )
    return StoredFile.objects.create(
        bucket=bucket, key=key, filename=filename,
        content_type=content_type, size=len(file_bytes),
    )


def _attach(client, stored, employee):
    root = get_or_create_root(client)
    folder = _mk(client, root, "Публикации ЕФРСБ", "efrsb", 6)
    ClientFile.objects.create(
        folder=folder, stored_file=stored, name=stored.filename,
        size=stored.size or 0, content_type=stored.content_type, uploaded_by=employee,
    )


def save_publication(publication, text: str, *, employee=None):
    """Сохранить итоговый текст: publication.generated_text + .docx/.pdf в папку клиента."""
    text = (text or "").strip()
    if not text:
        raise EfrsbGenError("Текст сообщения пуст — сгенерируйте или введите его.")

    title = publication.type_label
    docx_bytes = _text_to_docx(title, text)
    try:
        pdf_bytes = docx_to_pdf(docx_bytes)
    except Exception as exc:  # noqa: BLE001
        raise EfrsbGenError(f"Не удалось собрать PDF: {exc}") from exc

    base = f"ЕФРСБ — {title}"[:120]
    docx_sf = _store(docx_bytes, filename=f"{base}.docx", content_type=DOCX_CT)
    pdf_sf = _store(pdf_bytes, filename=f"{base}.pdf", content_type="application/pdf")

    client = publication.case.service.client
    _attach(client, pdf_sf, employee)
    _attach(client, docx_sf, employee)

    publication.title = title
    publication.generated_text = text
    publication.content_docx = docx_sf
    publication.content_pdf = pdf_sf
    publication.generated_at = timezone.now()
    if employee is not None and publication.created_by_id is None:
        publication.created_by = employee
    if publication.status == publication.STATUS_DRAFT:
        publication.status = publication.STATUS_GENERATED
    publication.save()

    try:
        client_log.invalidate_cache()
        client_log.record_action(
            client, "efrsb_text_generated",
            comment=f"Сформирован текст сообщения ЕФРСБ: {title}. "
                    f"Файл — в папке «Публикации ЕФРСБ».",
            employee=employee, stored_file=pdf_sf,
        )
    except Exception:
        log.exception("save_publication: не удалось записать событийку")
    return publication
