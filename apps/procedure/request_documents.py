"""Формирование документа-запроса (исходящее письмо): подстановка плейсхолдеров
из дела/должника/АУ/госоргана → .docx + PDF в файлы дела.

Переиспользует движок AFD (render_docx, docx_to_pdf, S3). Подпись/печать (PNG)
накладываются позже — когда заданы ArbitrationManager.signature_file/stamp_file.
"""
import logging
import re

from django.core.cache import cache
from django.db.models import Max
from django.utils import timezone

from apps.afd.docx_engine import render_docx
from apps.afd.pdf_utils import docx_to_pdf, pdf_page_count
from apps.crm import client_log
from apps.files.folder_utils import _mk, get_or_create_root
from apps.files.models import ClientFile, StoredFile
from apps.files.s3_utils import download_file_from_s3, upload_file_to_s3

from .recipient_resolver import RequestTypeLookup

log = logging.getLogger(__name__)
DOCX_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class RequestDocError(RuntimeError):
    pass


def _fmt(d):
    return d.strftime("%d.%m.%Y") if d else ""


def _spouse_data(spouse) -> str:
    if spouse is None:
        return ""
    fio = " ".join(filter(None, [spouse.last_name, spouse.first_name, spouse.patronymic])).strip()
    parts = []
    if spouse.birth_date:
        parts.append(f"дата рождения {_fmt(spouse.birth_date)}")
    if spouse.inn:
        parts.append(f"ИНН {spouse.inn}")
    if spouse.snils:
        parts.append(f"СНИЛС {spouse.snils}")
    return f"{fio} ({', '.join(parts)})" if parts else fio


def _debtor_address(client):
    addr = (client.addresses.filter(address_type="registration").first()
            or client.addresses.first())
    if not addr:
        return "", ""
    return (addr.postal_code or ""), (addr.result or addr.source or "")


# Ведущий почтовый индекс в строке адреса и приставка в названии суда.
_LEADING_INDEX_RE = re.compile(r"^\s*(\d{6})\s*,?\s*")
_COURT_PREFIX_RE = re.compile(r"^\s*(?:Арбитражный\s+суд|Арбитражного\s+суда|АС)\s+", re.IGNORECASE)


def _split_leading_index(idx: str, full: str) -> tuple[str, str]:
    """(индекс, адрес БЕЗ ведущего индекса) для шаблонов запросов.

    🛑 В шаблонах адрес печатается как «{индекс}, {адрес регистрации}», а DaData
    кладёт индекс ещё и в начало самого адреса — выходило «403342, 403342,
    Волгоградская обл, …». Режем именно здесь, а не в `_debtor_address`: её
    использует ещё и `apps.efrsb`, где `{адрес регистрации}` печатается один и
    индекс в нём нужен.
    """
    m = _LEADING_INDEX_RE.match(full or "")
    if not m:
        return idx, (full or "").strip()
    # Индекс не теряем: если в карточке его нет, берём из самой строки адреса.
    return (idx or m.group(1)), full[m.end():].strip()


def _court_short(service, arb) -> str:
    """Название суда БЕЗ слов «Арбитражный суд» — в шаблонах они уже напечатаны.

    🛑 Шаблоны говорят «Арбитражный суд {арбитражный суд}», а kad пишет в
    `court_name` «АС Волгоградской области» → получалось «Арбитражный суд АС
    Волгоградской области». Какой это суд, решаем тем же способом, что и ЕФРСБ
    (по префиксу номера дела «А12-…» → официальное название региона), иначе на
    делах, ушедших в апелляцию, в письмо попадала бы текущая инстанция вместе
    с фамилией судьи.
    """
    from apps.efrsb.generator import _court  # единый источник «какой это суд»
    try:
        name, _addr = _court(service, arb)
    except Exception:  # noqa: BLE001 — документ важнее красивого названия
        log.exception("Не удалось определить суд, берём название из kad")
        name = ""
    if not name:
        name = ((arb.court_name if arb is not None else "") or "").split("/")[0]
    return _COURT_PREFIX_RE.sub("", name).strip()


def _am_procedure(case):
    """Процедура, чей ФУ берём для реквизитов (актуальная с назначенным АУ)."""
    return (case.procedures.exclude(arbitr_manager=None).order_by("-order").first()
            or case.current_procedure
            or case.procedures.order_by("-order").first())


def build_request_context(req, *, marriage_cert="", gen_date=None) -> dict:
    case = req.case
    client = case.service.client
    proc = _am_procedure(case)
    am = proc.arbitr_manager if (proc and proc.arbitr_manager_id) else None
    rec = req.recipient
    arb = getattr(case.service, "arbitr_case", None)
    idx, addr_reg = _debtor_address(client)
    idx, addr_short = _split_leading_index(idx, addr_reg)
    gen_date = gen_date or timezone.localdate()
    rec_addr = ""
    if rec:
        rec_addr = rec.legal_address or rec.actual_address or rec.postal_address or ""
    elif (req.request_type_id
          and req.request_type.recipient_lookup == RequestTypeLookup.DEBTOR):
        # Уведомление должнику: адресат — сам клиент, адрес берём из его карточки.
        rec_addr = addr_reg
    return {
        # Должник
        "Фамилия": client.last_name or "", "Имя": client.first_name or "",
        "Отчество": client.patronymic or "",
        "дата рождения": _fmt(client.birth_date), "место рождения": client.birth_place or "",
        "СНИЛС": client.snils or "", "ИНН": client.inn or "",
        "индекс": idx, "адрес регистрации": addr_short,
        # Финуправляющий (АУ)
        "ФИО Финансовый управляющий": am.full_fio if am else "",
        "ФамилияИО АУ": am.short_fio if am else "",
        "ИНН АУ": am.inn if am else "", "СНИЛС АУ": am.snils if am else "",
        "Адрес арбитражного управляющего": am.corr_address if am else "",
        "Телефон арбитражного": am.phone if am else "", "email арбитражного": am.email if am else "",
        "Реквизиты СРО": am.sro_display if am else "",
        # Дело / суд
        "арбитражный суд": _court_short(case.service, arb),
        "номер дела": (arb.case_number if arb else ""),
        "дата решения": _fmt(proc.intro_date) if proc else "",
        "срок процедуры": (str(proc.term_months) if proc and proc.term_months else ""),
        # Запрос (исходящее)
        "Исх.№": str(req.outgoing_number) if req.outgoing_number else "",
        "Исх.дата": _fmt(gen_date),
        "Адресат": (rec.name if rec else (req.recipient_name or "")),
        "Адрес": rec_addr,
        # Супруг
        "данные на супруга": _spouse_data(client.spouse),
        "свидетельство о браке": marriage_cert or "",
    }


def _apply_signature(docx_bytes: bytes, am) -> bytes:
    """Вставить один PNG (подпись+печать вместе) в строку подписи ФУ — после
    «Финансовый управляющий», перед ФИО. Нет картинки — вернуть как есть."""
    import io
    sig = None
    if am and am.signature_file_id:
        try:
            sig = download_file_from_s3(am.signature_file.bucket, am.signature_file.key)
        except Exception:  # noqa: BLE001
            # Картинка недоступна (нет прав/файла в S3) — документ важнее подписи:
            # отдаём без неё, а не роняем всё формирование пакета.
            log.exception("Не удалось скачать подпись ФУ %s — документ без подписи", am.pk)
    if not sig:
        return docx_bytes
    from docx import Document
    from docx.shared import Cm
    doc = Document(io.BytesIO(docx_bytes))
    label = "Финансовый управляющий"
    target = None
    for p in doc.paragraphs:
        if label in p.text:
            target = p  # последняя строка подписи
    if target is None or not target.runs:
        return docx_bytes
    full = "".join(r.text for r in target.runs)
    idx = full.find(label)
    rest = full[idx + len(label):] if idx >= 0 else ""
    target.runs[0].text = label + " "
    for r in target.runs[1:]:
        r.text = ""
    target.add_run().add_picture(io.BytesIO(sig), width=Cm(5))
    target.add_run(" " + rest)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# Карта плейсхолдеров → (раздел, человекочитаемая метка) для предпроверки.
_PH_MAP = [
    ("Должник", [
        ("Фамилия", "Фамилия"), ("Имя", "Имя"), ("Отчество", "Отчество"),
        ("дата рождения", "Дата рождения"), ("место рождения", "Место рождения"),
        ("СНИЛС", "СНИЛС"), ("ИНН", "ИНН"),
        ("индекс", "Индекс (адрес)"), ("адрес регистрации", "Адрес регистрации"),
    ]),
    ("Финуправляющий (АУ)", [
        ("ФИО Финансовый управляющий", "ФИО ФУ"), ("ФамилияИО АУ", "Фамилия И.О. ФУ"),
        ("ИНН АУ", "ИНН ФУ"), ("СНИЛС АУ", "СНИЛС ФУ"),
        ("Адрес арбитражного управляющего", "Адрес корреспонденции ФУ"),
        ("Телефон арбитражного", "Телефон ФУ"), ("email арбитражного", "E-mail ФУ"),
        ("Реквизиты СРО", "СРО"),
    ]),
    ("Дело и суд", [
        ("арбитражный суд", "Арбитражный суд"), ("номер дела", "Номер дела"),
        ("дата решения", "Дата решения"), ("срок процедуры", "Срок процедуры, мес."),
    ]),
    ("Адресат (госорган)", [
        ("Адресат", "Адресат"), ("Адрес", "Адрес госоргана"),
    ]),
    ("Супруг", [
        ("данные на супруга", "Данные супруга"),
        ("свидетельство о браке", "Свидетельство о браке"),
    ]),
    ("Исходящее", [
        ("Исх.№", "Исходящий №"), ("Исх.дата", "Дата исходящего"),
    ]),
]
_AUTO_KEYS = {"Исх.№", "Исх.дата"}  # присваиваются автоматически при формировании


def template_placeholders(tpl):
    """Плейсхолдеры .docx-шаблона (кэш в Redis по StoredFile).

    Нужен кэш: модалка пакета проверяет готовность сразу ~18 типов, а качать
    столько .docx из S3 на каждое открытие — секунды ожидания. Шаблон меняется
    редко и всегда через новый StoredFile → ключ по его id безопасен.
    """
    if not (tpl and tpl.stored_file_id):
        return None
    ckey = f"afd:tpl_ph:{tpl.stored_file_id}"
    cached = cache.get(ckey)
    if cached is not None:
        return set(cached)
    try:
        from apps.afd.docx_engine import list_placeholders
        tb = download_file_from_s3(tpl.stored_file.bucket, tpl.stored_file.key)
        found = sorted(list_placeholders(tb))
    except Exception:
        log.exception("template_placeholders: не удалось прочитать шаблон %s", tpl.pk)
        return None
    cache.set(ckey, found, 60 * 60 * 24)
    return set(found)


def _check_key(ctx, key):
    """Одно поле: заполнено ли и в верном ли формате."""
    val = (ctx.get(key) or "").strip()
    if key in _AUTO_KEYS:
        return {"value": "присвоится автоматически", "ok": True, "note": ""}
    if key == "свидетельство о браке":
        return {"value": "", "ok": True, "note": "вводится в форме ниже"}
    if not val:
        return {"value": "", "ok": False, "note": "не заполнено"}
    if key in ("ИНН", "ИНН АУ") and len(re.sub(r"\D", "", val)) not in (10, 12):
        return {"value": val, "ok": False, "note": "неверный формат ИНН"}
    if key in ("СНИЛС", "СНИЛС АУ") and len(re.sub(r"\D", "", val)) != 11:
        return {"value": val, "ok": False, "note": "неверный формат СНИЛС"}
    if key == "email арбитражного" and "@" not in val:
        return {"value": val, "ok": False, "note": "неверный e-mail"}
    return {"value": val, "ok": True, "note": ""}


def check_request_data(req):
    """Предпроверка данных для подстановки: какие плейсхолдеры шаблона заполнены
    (ok) и каких нет / неверный формат. Возвращает (all_ok, groups)."""
    ctx = build_request_context(req)
    used = set(ctx.keys())
    tpl = req.request_type.template if req.request_type_id else None
    found = template_placeholders(tpl)
    if found:
        used = found

    def _check(key):
        return _check_key(ctx, key)

    groups, all_ok = [], True
    known = set()
    for gname, items in _PH_MAP:
        rows = []
        for key, label in items:
            known.add(key)
            if key not in used:
                continue
            chk = _check(key)
            chk["label"] = label
            rows.append(chk)
            all_ok = all_ok and chk["ok"]
        if rows:
            groups.append({"name": gname, "rows": rows})
    extra = sorted(k for k in used if k not in known)
    if extra:
        rows = []
        for key in extra:
            chk = _check(key)
            chk["label"] = key
            rows.append(chk)
            all_ok = all_ok and chk["ok"]
        groups.append({"name": "Прочее", "rows": rows})
    return all_ok, groups


# Группы плейсхолдеров, не зависящие от типа запроса, — общие сведения дела.
_CASE_GROUPS = ("Должник", "Финуправляющий (АУ)", "Дело и суд")
# Опциональные группы: их пустота не мешает сформировать документ.
_OPTIONAL_GROUPS = ("Супруг", "Исходящее")


def _probe_request(case, rt=None, recipient=None):
    """Несохранённый Request — чтобы прогнать build_request_context до создания."""
    from . import services
    from .models import Request
    from .recipient_resolver import RequestTypeLookup

    name = ""
    if recipient is not None:
        name = recipient.short_name or recipient.name
    elif rt is not None and rt.recipient_lookup == RequestTypeLookup.DEBTOR:
        name = services.debtor_display(case.service.client)
    return Request(case=case, request_type=rt, recipient=recipient, recipient_name=name)


def check_case_data(case):
    """Общие сведения дела: должник / финуправляющий / дело и суд.

    Возвращает (all_ok, groups) — в groups ТОЛЬКО пробелы (для предупреждения
    в шапке модалки пакета). Плюс отдельно ловим отсутствие самого ФУ и его
    подписи: без них документы формируются, но без блока подписи.
    """
    ctx = build_request_context(_probe_request(case))
    groups, all_ok = [], True
    for gname, items in _PH_MAP:
        if gname not in _CASE_GROUPS:
            continue
        rows = []
        for key, label in items:
            chk = _check_key(ctx, key)
            if chk["ok"]:
                continue
            chk["label"] = label
            rows.append(chk)
            all_ok = False
        if rows:
            groups.append({"name": gname, "rows": rows})

    proc = _am_procedure(case)
    am = proc.arbitr_manager if (proc and proc.arbitr_manager_id) else None
    extra = []
    if am is None:
        extra.append({"label": "Финансовый управляющий",
                      "value": "", "ok": False,
                      "note": "не назначен в процедуре — реквизиты ФУ будут пустыми"})
    elif not am.signature_file_id:
        extra.append({"label": "Подпись и печать ФУ",
                      "value": "", "ok": False,
                      "note": "нет файла подписи — документ сформируется без неё"})
    if extra:
        all_ok = False
        groups.append({"name": "Финуправляющий (АУ)", "rows": extra})
    return all_ok, groups


def check_request_ready(case, rt, recipient, *, creditors_count=None):
    """Готовность ОДНОЙ позиции пакета: (ready, issues).

    ready — можно сформировать документ: есть шаблон, определён адресат и
    заполнены все плейсхолдеры шаблона. Данные супруга и исходящий № не в счёт
    (первое опционально, второе присваивается автоматически).
    """
    from .recipient_resolver import RequestTypeLookup

    lookup = rt.recipient_lookup
    issues = []

    tpl = rt.template if rt.template_id else None
    if tpl is None:
        issues.append("нет шаблона документа")

    needs_recipient = lookup in (
        RequestTypeLookup.REGION, RequestTypeLookup.FNS, RequestTypeLookup.MANUAL)
    no_recipient = needs_recipient and recipient is None
    if no_recipient:
        issues.append("не выбран адресат")
    if lookup == RequestTypeLookup.CREDITORS and creditors_count == 0:
        issues.append("в анкете нет кредиторов")

    ctx = build_request_context(_probe_request(case, rt, recipient))
    used = template_placeholders(tpl)
    if used is None:
        used = set(ctx.keys())

    # Адресат в {Адресат}/{Адрес} проверяем, только когда он вообще должен быть
    # в документе: СМЭВ-заглушке адресат не нужен; письмам кредиторам он
    # подставится при создании (по письму на каждого); а если адресата ещё не
    # выбрали — про это уже сказано отдельной строкой, дублировать не надо.
    skip_recipient_keys = (
        no_recipient
        or lookup in (RequestTypeLookup.NONE, RequestTypeLookup.CREDITORS)
    )
    missing = []
    for gname, items in _PH_MAP:
        if gname in _OPTIONAL_GROUPS:
            continue
        for key, label in items:
            if key not in used:
                continue
            if skip_recipient_keys and key in ("Адресат", "Адрес"):
                continue
            if not _check_key(ctx, key)["ok"]:
                missing.append(label)
    if missing:
        issues.append("не заполнено: " + ", ".join(missing))
    return (not issues), issues


def _has_image(p):
    from docx.oxml.ns import qn
    return bool(p._element.findall(".//" + qn("w:drawing")))


def extract_editable_paragraphs(docx_bytes: bytes) -> list:
    """Абзацы документа для редактирования (непустые, без картинок).
    index — позиция в стабильном обходе (тот же, что в apply)."""
    import io
    from docx import Document
    from apps.afd.docx_engine import _iter_paragraphs
    doc = Document(io.BytesIO(docx_bytes))
    out = []
    for i, p in enumerate(_iter_paragraphs(doc)):
        if _has_image(p):
            continue
        if p.text.strip():
            out.append({"index": i, "text": p.text})
    return out


def apply_paragraph_edits(docx_bytes: bytes, edits: dict) -> bytes:
    """Применить правки текста абзацев (edits: {index: new_text}). Текст пишется
    в первый run абзаца (его формат сохраняется), остальные очищаются. Абзацы с
    картинками не трогаются."""
    import io
    from docx import Document
    from apps.afd.docx_engine import _iter_paragraphs
    doc = Document(io.BytesIO(docx_bytes))
    for i, p in enumerate(_iter_paragraphs(doc)):
        if i not in edits or _has_image(p):
            continue
        new_text = edits[i]
        if p.runs:
            p.runs[0].text = new_text
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.add_run(new_text)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def save_edited_document(req, new_docx_bytes: bytes, *, employee=None):
    """Пересохранить отредактированный .docx: re-PDF + S3 + подшивка + обновить req."""
    pdf_bytes = docx_to_pdf(new_docx_bytes)
    base = request_document_basename(req)
    docx_sf = _store(new_docx_bytes, filename=f"{base}.docx", content_type=DOCX_CT)
    pdf_sf = _store(pdf_bytes, filename=f"{base}.pdf", content_type="application/pdf")
    client = req.case.service.client
    _attach(client, pdf_sf, employee)
    _attach(client, docx_sf, employee)
    req.document_docx = docx_sf
    req.document_pdf = pdf_sf
    req.pages_count = pdf_page_count(pdf_bytes) or None
    req.generated_at = timezone.now()
    req.pdf_built_at = req.generated_at
    req.save(update_fields=[
        "document_docx", "document_pdf", "pages_count", "generated_at",
        "pdf_built_at", "updated_at",
    ])
    return req


def rebuild_pdf_from_docx(req):
    """Пересобрать PDF из текущего .docx запроса (кнопка «Пересобрать PDF»).

    Нужна после правки в онлайн-редакторе: Collabora меняет только .docx, а
    когда собирать PDF — решает юрист. Автоматически на каждое автосохранение
    гонять LibreOffice бессмысленно (секунды CPU на каждое нажатие).

    🛑 Перезаписываем ТОТ ЖЕ объект в S3, а не заводим новый StoredFile: иначе
    в папке «Запросы» файл-менеджера копилась бы куча версий одного письма, а
    ссылки на PDF в открытых вкладках протухали бы.
    """
    from apps.files.s3_utils import upload_file_to_s3_key

    if not req.document_docx_id:
        raise RequestDocError("Нет .docx — из чего собирать PDF.")
    docx_sf = req.document_docx
    docx_bytes = download_file_from_s3(docx_sf.bucket, docx_sf.key)
    pdf_bytes = docx_to_pdf(docx_bytes)

    pdf_sf = req.document_pdf
    if pdf_sf is not None:
        upload_file_to_s3_key(pdf_bytes, bucket=pdf_sf.bucket, key=pdf_sf.key,
                              content_type="application/pdf")
        pdf_sf.size = len(pdf_bytes)
        pdf_sf.save(update_fields=["size"])
    else:
        # PDF-двойника не было (документ подгружали одним .docx) — заводим.
        base = request_document_basename(req)
        req.document_pdf = _store(pdf_bytes, filename=f"{base}.pdf",
                                  content_type="application/pdf")
    req.pages_count = pdf_page_count(pdf_bytes) or None
    req.pdf_built_at = timezone.now()
    req.save(update_fields=["document_pdf", "pages_count", "pdf_built_at", "updated_at"])
    return req


# ── Понятное имя файла документа запроса («Запрос в ЛРР по Каныгину ДВ») ─────
import re as _re

_FN_BAD = _re.compile(r'[\\/:*?"<>|\r\n\t]+')


def _safe_filename(name: str) -> str:
    return _re.sub(r"\s{2,}", " ", _FN_BAD.sub(" ", name or "")).strip()


def _guess_gender(client) -> str:
    g = (getattr(client, "gender", "") or "").strip().lower()
    if g in ("male", "female"):
        return g
    pa = (getattr(client, "patronymic", "") or "").strip().lower()
    if pa.endswith(("вна", "чна", "нична")):
        return "female"
    if pa.endswith(("вич", "ич", "ыч")):
        return "male"
    return ""


def _surname_dative(last: str, gender: str) -> str:
    """Грубая дательная форма фамилии для имени файла (не уверены → как есть)."""
    l = (last or "").strip()
    if not l:
        return l
    low = l.lower().replace("ё", "е")
    # Явные женские формы фамилий (по окончанию, независимо от пола-поля).
    if low.endswith(("ова", "ева", "ина", "ына")):
        return l[:-1] + "ой"
    if low.endswith(("ская", "цкая")):
        return l[:-2] + "ой"
    # Явные мужские формы.
    if low.endswith(("ов", "ев", "ин", "ын", "цын")):
        return l + "у"
    if low.endswith(("ский", "цкий", "ской", "цкой")):
        return l[:-2] + "ому"
    # Несклоняемые (-ых/-их/-ко/-аго, гласный на конце кроме -а/-я).
    if low.endswith(("ых", "их", "ко", "аго", "яго", "ово")) or low[-1] in "оеиуюыэ":
        return l
    # Фамилии-прилагательные.
    if gender == "female" and low.endswith(("ая", "яя")):
        return l[:-2] + "ой"
    if gender != "female" and low.endswith(("ый", "ой")):
        return l[:-2] + "ому"
    if gender != "female" and low.endswith("ий"):
        return l[:-2] + "ему"
    # Согласная / мягкий знак — склоняется только мужская.
    if low.endswith("ь"):
        return (l[:-1] + "ю") if gender != "female" else l
    if low[-1] in "бвгдзйклмнпрстфхцчшщ":
        return (l + "у") if gender != "female" else l
    return l


def _client_dative_short(client) -> str:
    """«Каныгину ДВ» — фамилия в дательном + инициалы без точек (как просил юрист)."""
    if client is None:
        return ""
    fi = (getattr(client, "first_name", "") or "").strip()
    pa = (getattr(client, "patronymic", "") or "").strip()
    initials = (fi[:1].upper() if fi else "") + (pa[:1].upper() if pa else "")
    surn = _surname_dative(getattr(client, "last_name", "") or "", _guess_gender(client))
    return (surn + (" " + initials if initials else "")).strip()


def request_document_basename(req) -> str:
    """Понятное имя файла: «<тип без скобок> по <Фамилия Инициалы>»."""
    title = (getattr(req, "title", "") or "Запрос").split("(")[0].strip()
    try:
        client = req.case.service.client
    except Exception:
        client = None
    who = _client_dative_short(client)
    name = f"{title} по {who}".strip() if who else title
    return _safe_filename(name)[:120] or "Запрос"


def _store(file_bytes, *, filename, content_type):
    bucket, key = upload_file_to_s3(
        file_bytes, prefix="procedure/requests", filename=filename, content_type=content_type,
    )
    return StoredFile.objects.create(
        bucket=bucket, key=key, filename=filename, content_type=content_type, size=len(file_bytes),
    )


def _attach(client, stored, employee):
    root = get_or_create_root(client)
    folder = _mk(client, root, "Запросы", "requests", 5)
    ClientFile.objects.create(
        folder=folder, stored_file=stored, name=stored.filename,
        size=stored.size or 0, content_type=stored.content_type, uploaded_by=employee,
    )


def generate_request_document(req, *, with_signature=False, marriage_cert="", employee=None):
    """Сформировать документ запроса. Возвращает req (с document_pdf/docx).

    🛑 Наложение подписи/печати (PNG) — TODO, когда заданы signature_file/stamp_file
    у АУ; сейчас `with_signature` только сохраняется.
    """
    rtype = req.request_type
    tpl = rtype.template if rtype else None
    if tpl is None or not tpl.stored_file_id:
        raise RequestDocError(
            "У типа запроса не задан шаблон документа. Привяжите .docx в справочнике «Типы запросов»."
        )
    # Исходящий номер (сквозной по делу) — присваиваем один раз.
    if not req.outgoing_number:
        mx = req.case.requests.aggregate(m=Max("outgoing_number"))["m"] or 0
        req.outgoing_number = mx + 1

    ctx = build_request_context(req, marriage_cert=marriage_cert)
    try:
        template_bytes = download_file_from_s3(tpl.stored_file.bucket, tpl.stored_file.key)
    except Exception as exc:
        # Типичный случай на dev: шаблон залит в прод-бакет, у dev-ключа туда нет
        # прав. Раньше это выглядело как «внутренняя ошибка» и требовало логов.
        log.exception("Шаблон %s недоступен в S3", tpl.pk)
        raise RequestDocError(
            f"Файл шаблона «{tpl.name}» недоступен в хранилище "
            f"(бакет {tpl.stored_file.bucket}). Перезалейте шаблон в справочнике "
            f"«Типы запросов» или выполните `manage.py load_request_templates --force`."
        ) from exc
    docx_bytes = render_docx(template_bytes, ctx)
    if with_signature:
        proc = _am_procedure(req.case)
        am = proc.arbitr_manager if (proc and proc.arbitr_manager_id) else None
        if am is not None:
            docx_bytes = _apply_signature(docx_bytes, am)
    pdf_bytes = docx_to_pdf(docx_bytes)

    base = request_document_basename(req)
    docx_sf = _store(docx_bytes, filename=f"{base}.docx", content_type=DOCX_CT)
    pdf_sf = _store(pdf_bytes, filename=f"{base}.pdf", content_type="application/pdf")

    client = req.case.service.client
    _attach(client, pdf_sf, employee)
    _attach(client, docx_sf, employee)

    req.document_pdf = pdf_sf
    req.document_docx = docx_sf
    req.pages_count = pdf_page_count(pdf_bytes) or None
    req.with_signature = bool(with_signature)
    req.generated_at = timezone.now()
    req.pdf_built_at = req.generated_at  # собраны вместе — PDF актуален
    req.save(update_fields=[
        "outgoing_number", "document_pdf", "document_docx", "pages_count",
        "with_signature", "generated_at", "pdf_built_at", "updated_at",
    ])
    try:
        from apps.crm.models import ActionType
        ActionType.objects.get_or_create(
            code="request_document_created",
            defaults={"name": "Сформирован запрос в госорган", "order": 36, "is_manual": False},
        )
        client_log.invalidate_cache()
        client_log.record_action(
            client, "request_document_created",
            comment=f"Сформирован запрос (исх. № {req.outgoing_number}): "
                    f"{req.title} → {req.recipient_display}. Файл — в папке «Запросы».",
            employee=employee, stored_file=pdf_sf,
        )
    except Exception:
        log.exception("generate_request_document: не удалось записать событийку")
    return req
