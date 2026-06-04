# -*- coding: utf-8 -*-
"""Приложения к заявлению о банкротстве (python-docx):

- creditors_form_docx  — «Список кредиторов и должников гражданина» (форма)
- property_form_docx    — «Опись имущества гражданина» (форма)
- petition_docx         — «Ходатайство о введении процедуры реализации имущества»

Возвращают bytes .docx. Данные берутся из ctx (isk_context.build_isk_context)
и из списка creditors / overrides.
"""
import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from .isk_engine import FONT, SIZE, _base_doc, add_para


def _doc():
    return _base_doc()


def _p(doc, text, align="both", bold=False, size=SIZE, indent=False):
    return add_para(doc, text, align=align, bold=bold, size=size, indent=indent)


def _set_cell(cell, text, bold=False, size=10, align="left"):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                      "both": WD_ALIGN_PARAGRAPH.JUSTIFY}.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    para.paragraph_format.space_after = Pt(0)
    r = para.add_run(str(text) if text is not None else "")
    r.bold = bold
    r.font.name = FONT
    r.font.size = Pt(size)


def _table(doc, headers, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, (h, w) in enumerate(zip(headers, widths)):
        _set_cell(t.rows[0].cells[i], h, bold=True, size=9, align="center")
        t.columns[i].width = Cm(w)
    return t


# ── Информация о гражданине (общий блок форм) ───────────────────────────────
def _debtor_info(doc, ctx):
    rows = [
        ("Фамилия", ctx.get("debtor_last", "")),
        ("Имя", ctx.get("debtor_first", "")),
        ("Отчество", ctx.get("debtor_patronymic", "")),
        ("Прежние Ф.И.О. (при наличии)", ctx.get("former_name", "")),
        ("Дата рождения", ctx.get("birth_date", "")),
        ("Место рождения", ctx.get("birth_place", "")),
        ("СНИЛС", ctx.get("snils", "")),
        ("ИНН (при наличии)", ctx.get("inn", "")),
        ("Документ, удостоверяющий личность",
         f"Паспорт, серия {ctx.get('passport_series','')} № {ctx.get('passport_number','')}, "
         f"выдан {ctx.get('passport_issued_by','')} {ctx.get('passport_issued_date','')}, "
         f"код подразделения {ctx.get('passport_division_code','')}"),
        ("Адрес регистрации по месту жительства", ctx.get("reg_address", "")),
    ]
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    for label, val in rows:
        cells = t.add_row().cells
        _set_cell(cells[0], label, size=9)
        _set_cell(cells[1], val, size=9)
    t.columns[0].width = Cm(6)
    t.columns[1].width = Cm(11)


_KIND_CONTENT = {
    "bank": "Кредит", "mfo": "Заём", "utility": "Коммунальные услуги",
    "court": "По решению суда", "fine": "Штраф", "other": "Иное обязательство",
}


def creditors_form_docx(ctx, creditors) -> bytes:
    from .isk_context import fmt_money
    doc = _doc()
    _p(doc, "Список кредиторов и должников гражданина", align="center", bold=True)
    _p(doc, "Информация о гражданине", bold=True)
    _debtor_info(doc, ctx)
    _p(doc, "")
    _p(doc, "I. Сведения о кредиторах гражданина (по денежным обязательствам и (или) "
            "обязанности по уплате обязательных платежей)", bold=True, size=10)
    _p(doc, "1. Денежные обязательства", bold=True, size=10)
    t = _table(doc,
               ["N", "Содержание обязательства", "Кредитор",
                "Место нахождения кредитора", "Основание возникновения",
                "Сумма обязательства, всего", "в том числе задолженность",
                "Штрафы, пени"],
               [0.8, 2.5, 3.5, 4.0, 3.5, 2.6, 2.6, 1.8])
    for i, c in enumerate(creditors, 1):
        cells = t.add_row().cells
        basis = c.get("basis", "")
        if c.get("date"):
            basis += f" от {c['date']} г."
        amt = fmt_money(c.get("amount"))
        vals = [str(i), _KIND_CONTENT.get(c.get("kind"), "Обязательство"),
                c.get("name", ""), c.get("address", ""), basis, amt, amt, ""]
        for j, v in enumerate(vals):
            _set_cell(cells[j], v, size=9)
    _p(doc, "")
    _p(doc, "2. Обязательные платежи (налоги, сборы, иные обязательные платежи): "
            "сведения отсутствуют / указываются при наличии.", size=9)
    _p(doc, "")
    _p(doc, "II. Сведения о кредиторах по обязательствам из предпринимательской "
            "деятельности: отсутствуют.", size=9)
    _p(doc, "III. Сведения о должниках гражданина: отсутствуют.", size=9)
    _p(doc, "")
    _p(doc, f"«___» __________ {ctx.get('year','')} г.\t\t_____________ /{ctx.get('debtor_short','')}")
    out = io.BytesIO(); doc.save(out); return out.getvalue()


def property_form_docx(ctx, overrides=None) -> bytes:
    overrides = overrides or {}
    doc = _doc()
    _p(doc, "Опись имущества гражданина", align="center", bold=True)
    _p(doc, "Информация о гражданине", bold=True)
    _debtor_info(doc, ctx)
    _p(doc, "")

    # I. Недвижимое имущество
    _p(doc, "I. Недвижимое имущество", bold=True, size=10)
    t1 = _table(doc, ["N", "Вид и наименование имущества", "Вид собственности",
                      "Местонахождение (адрес)", "Площадь (кв. м)",
                      "Основание приобретения и стоимость", "Залог"],
                [0.8, 3.5, 2.5, 4.0, 2.0, 3.5, 2.0])
    realty = overrides.get("property_realty") or []
    if not realty:
        _set_cell(t1.add_row().cells[1], "—", size=9)
    for i, r in enumerate(realty, 1):
        cells = t1.add_row().cells
        for j, v in enumerate([str(i), r.get("name", ""), r.get("ownership", ""),
                               r.get("address", ""), r.get("area", ""),
                               r.get("basis", ""), r.get("pledge", "")]):
            _set_cell(cells[j], v, size=9)
    _p(doc, "")

    # II. Движимое имущество
    _p(doc, "II. Движимое имущество (транспортные средства)", bold=True, size=10)
    t2 = _table(doc, ["N", "Вид, марка, модель, год", "VIN/идент. номер",
                      "Вид собственности", "Местонахождение", "Стоимость", "Залог"],
                [0.8, 3.5, 3.0, 2.5, 3.5, 2.0, 2.0])
    movable = overrides.get("property_movable") or []
    if not movable:
        _set_cell(t2.add_row().cells[1], "—", size=9)
    for i, r in enumerate(movable, 1):
        cells = t2.add_row().cells
        for j, v in enumerate([str(i), r.get("name", ""), r.get("vin", ""),
                               r.get("ownership", ""), r.get("address", ""),
                               r.get("value", ""), r.get("pledge", "")]):
            _set_cell(cells[j], v, size=9)
    _p(doc, "")

    # III. Счета в банках
    _p(doc, "III. Сведения о счетах в банках и иных кредитных организациях", bold=True, size=10)
    t3 = _table(doc, ["N", "Наименование и адрес банка", "Вид и валюта счёта", "Остаток (руб.)"],
                [0.8, 8.0, 4.0, 4.0])
    accounts = overrides.get("property_accounts") or []
    if not accounts:
        _set_cell(t3.add_row().cells[1], "—", size=9)
    for i, r in enumerate(accounts, 1):
        cells = t3.add_row().cells
        for j, v in enumerate([str(i), r.get("bank", ""), r.get("type", "Текущий (рубль)"),
                               r.get("balance", "0,00")]):
            _set_cell(cells[j], v, size=9)
    _p(doc, "")
    _p(doc, "IV. Акции и иное участие в коммерческих организациях: отсутствуют.", size=9)
    _p(doc, "V. Иные ценные бумаги: отсутствуют.", size=9)
    _p(doc, "VI. Наличные денежные средства и иное ценное имущество: "
            + (overrides.get("property_cash") or "отсутствуют") + ".", size=9)
    _p(doc, "")
    _p(doc, f"«___» __________ {ctx.get('year','')} г.\t\t_____________ /{ctx.get('debtor_short','')}")
    out = io.BytesIO(); doc.save(out); return out.getvalue()


def petition_docx(ctx) -> bytes:
    doc = _doc()
    _p(doc, f"В {ctx.get('court_name','')}")
    if ctx.get("court_address"):
        _p(doc, ctx["court_address"])
    _p(doc, f"Должник (заявитель): {ctx.get('debtor_full','')} "
            f"(ИНН: {ctx.get('inn','')} СНИЛС: {ctx.get('snils','')}), "
            f"адрес: {ctx.get('reg_address','')}")
    _p(doc, "")
    _p(doc, "ХОДАТАЙСТВО", align="center", bold=True)
    _p(doc, "о введении процедуры реализации имущества", align="center", bold=True)
    _p(doc, "")
    _p(doc, "Согласно ст. 213.6 п. 8 Федерального закона № 127-ФЗ от 26.10.2002 г.: «По "
            "результатам рассмотрения обоснованности заявления о признании гражданина "
            "банкротом, если гражданин не соответствует требованиям для утверждения плана "
            "реструктуризации долгов, установленным пунктом 1 статьи 213.13 настоящего "
            "Федерального закона, арбитражный суд вправе на основании ходатайства "
            "гражданина вынести решение о признании его банкротом и введении процедуры "
            "реализации имущества гражданина».", indent=True)
    _p(doc, "В силу п. 30 Постановления Пленума Верховного Суда РФ от 13.10.2015 № 45 суд "
            "утверждает план реструктуризации долгов только в случае, если он одобрен "
            "должником. Отсутствие у должника достаточного дохода не предполагает с "
            "разумной долей вероятности возможности исполнить денежные обязательства по "
            "погашению кредиторской задолженности на условиях её отсрочки (рассрочки), что "
            "свидетельствует о невыполнимости какого-либо плана реструктуризации долгов "
            "гражданина и нецелесообразности введения данной процедуры.", indent=True)
    _p(doc, "На основании вышеизложенного, ходатайствую о вынесении судом решения о "
            "признании Должника несостоятельным (банкротом) с введением процедуры "
            "реализации имущества гражданина.", indent=True)
    _p(doc, "")
    _p(doc, f"«___» __________ {ctx.get('year','')} г.\t\t_____________ /{ctx.get('debtor_short','')}")
    out = io.BytesIO(); doc.save(out); return out.getvalue()
