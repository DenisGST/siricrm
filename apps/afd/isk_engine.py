# -*- coding: utf-8 -*-
"""Рендер заявления о банкротстве из секций (python-docx).

render_isk_docx(template, ctx, flags) -> bytes
  template — IskTemplate; секции рендерятся по порядку, условные пропускаются.
  ctx — плоский dict плейсхолдеров; flags — булевы флаги для include_condition.
"""
import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

_PH = re.compile(r"\{[^{}\n]+\}")
_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
    "both": WD_ALIGN_PARAGRAPH.JUSTIFY, "right": WD_ALIGN_PARAGRAPH.RIGHT,
}
FONT = "Times New Roman"
SIZE = 12


def substitute(text, ctx):
    def _s(m):
        key = m.group(0)[1:-1]
        if key in ctx:
            v = ctx[key]
            return "" if v is None else str(v)
        return m.group(0)  # неизвестный плейсхолдер оставляем видимым
    return _PH.sub(_s, text or "")


def _base_doc():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(SIZE)
    try:
        st.element.rPr.rFonts.set(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", FONT)
    except Exception:
        pass
    for s in doc.sections:
        s.top_margin = Cm(2); s.bottom_margin = Cm(2)
        s.left_margin = Cm(3); s.right_margin = Cm(1.5)
    return doc


def add_para(doc, text, *, align="both", bold=False, indent=False, size=SIZE,
             left_indent_cm=None):
    p = doc.add_paragraph()
    p.alignment = _ALIGN.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    pf = p.paragraph_format
    pf.space_after = Pt(0)
    pf.line_spacing = 1.15
    if indent:
        pf.first_line_indent = Cm(1.25)
    if left_indent_cm:
        pf.left_indent = Cm(left_indent_cm)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = FONT
    r.font.size = Pt(size)
    return p


def render_isk_docx(template, ctx, flags) -> bytes:
    doc = _base_doc()
    sections = template.sections.filter(is_active=True).order_by("order")
    for sec in sections:
        if sec.include_condition and not flags.get(sec.include_condition):
            continue
        text = substitute(sec.body, ctx)
        # «Шапка» (суд/должник/кредиторы) — смещается в правую половину листа
        # по правилам делопроизводства: большой левый отступ + выравнивание влево.
        is_header = sec.block_type in ("court_header", "creditors_header")
        line_align = "left" if is_header else sec.align
        left_cm = 8 if is_header else None
        indent = (not is_header) and sec.block_type in ("text", "petition") \
            and sec.align == "both"
        if sec.title:
            add_para(doc, sec.title, align="center", bold=True)
        lines = text.split("\n")
        for line in lines:
            if line.strip() == "":
                add_para(doc, "")  # пустая строка-разделитель
                continue
            add_para(doc, line, align=line_align, bold=sec.bold, indent=indent,
                     left_indent_cm=left_cm)
        # отступ между секциями
        add_para(doc, "")
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
