"""Формирование бланка заявки-договора на публикацию в «Коммерсантъ».

Что делаем: берём официальный бланк ИД (ФЛ), подставляем данные дела, отмечаем
чекбоксы, накладываем факсимиле АУ, конвертируем в PDF и подшиваем оба файла
в папку клиента «Публикации Коммерсантъ».

🛑 Шаблон — `OLD/Коммерсант/Заявка Коммерсантъ (ФЛ) — шаблон.docx`, получен из
   оригинала `blank-person.docx` с сайта ИД: Word-контролы (<w:sdt>) развёрнуты в
   обычные run'ы, а тексты-заглушки заменены на плейсхолдеры {…}. Без этого
   render_docx поля не видит — python-docx не отдаёт run'ы внутри content control.
   Если ИД поменяет бланк — перегенерировать шаблон, а не править этот модуль.

🛑 В письмо уходят ОБА файла: подписанный PDF и .docx — этого требует регламент
   приёма заявок по e-mail (bankruptcy.kommersant.ru/index.php?publemail).
"""
from __future__ import annotations

import io
import logging
from pathlib import Path

from django.conf import settings
from django.utils import timezone

from apps.afd.docx_engine import render_docx
from apps.afd.pdf_utils import docx_to_pdf
from apps.crm import client_log
from apps.efrsb.generator import _procedure_for
from apps.files.folder_utils import _mk, get_or_create_root
from apps.files.models import ClientFile, StoredFile
from apps.files.s3_utils import download_file_from_s3, upload_file_to_s3

from .generator import build_context
from .models import KommersantMessageType

log = logging.getLogger(__name__)
DOCX_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

BLANK_TEMPLATE = (
    Path(settings.BASE_DIR) / "OLD" / "Коммерсант" / "Заявка Коммерсантъ (ФЛ) — шаблон.docx"
)

CHECKED = "☒"
UNCHECKED = "☐"

# Поля бланка, которые обязан заполнить ИД-регламент (для предпроверки перед отправкой).
BLANK_REQUIRED_KEYS = [
    "ФИО Финансовый управляющий", "Реквизиты СРО", "Адрес арбитражного управляющего",
    "СНИЛС АУ", "ИНН АУ", "контакты АУ",
    "ФИО должника", "дата рождения", "место рождения", "адрес регистрации",
    "ИНН", "СНИЛС",
    "арбитражный суд", "номер дела", "дата решения",
]


class KommersantBlankError(RuntimeError):
    pass


def _load_template() -> bytes:
    if not BLANK_TEMPLATE.exists():
        raise KommersantBlankError(
            f"Не найден шаблон бланка заявки: {BLANK_TEMPLATE}. "
            f"Он должен лежать в репозитории (папка OLD/Коммерсант)."
        )
    return BLANK_TEMPLATE.read_bytes()


def build_blank_context(publication) -> dict:
    """Контекст бланка: данные дела + отметки чекбоксов + текст сообщения."""
    case = publication.case
    mt = publication.message_type
    ctx = build_context(
        case, message_type=mt, procedure=publication.procedure,
        overrides=publication.overrides,
    )

    checkbox = mt.blank_checkbox if mt else ""
    ctx.update({
        "текст сообщения": publication.text or "",
        "чек реструктуризация": (
            CHECKED if checkbox == KommersantMessageType.CHECKBOX_RESTRUCTURING else UNCHECKED
        ),
        "чек реализация": (
            CHECKED if checkbox == KommersantMessageType.CHECKBOX_REALIZATION else UNCHECKED
        ),
        "чек отчетные АУ": (
            CHECKED if publication.accounting_docs_to == publication.DOCS_TO_MANAGER else UNCHECKED
        ),
        "чек отчетные должник": (
            CHECKED if publication.accounting_docs_to == publication.DOCS_TO_DEBTOR else UNCHECKED
        ),
    })
    return ctx


def _apply_signature(docx_bytes: bytes, am) -> bytes:
    """Вставить факсимиле АУ в строку подписи бланка.

    Строка подписи — последняя таблица бланка: [ФИО] [ ] [подпись] [ ].
    Картинку кладём во вторую ячейку — над словом «подпись». Нет картинки или
    файл недоступен — возвращаем документ как есть: заявку лучше отдать без
    факсимиле (юрист подпишет вручную), чем уронить формирование.
    """
    sig = None
    if am is not None and am.signature_file_id:
        try:
            sig = download_file_from_s3(am.signature_file.bucket, am.signature_file.key)
        except Exception:  # noqa: BLE001
            log.exception("Коммерсантъ: не скачалась подпись АУ %s — бланк без факсимиле", am.pk)
    if not sig:
        return docx_bytes

    from docx import Document
    from docx.shared import Cm

    doc = Document(io.BytesIO(docx_bytes))
    if not doc.tables:
        return docx_bytes
    table = doc.tables[-1]
    if not table.rows or len(table.rows[0].cells) < 2:
        return docx_bytes

    cell = table.rows[0].cells[1]
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    para.add_run().add_picture(io.BytesIO(sig), width=Cm(4))
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _store(file_bytes: bytes, *, filename: str, content_type: str) -> StoredFile:
    bucket, key = upload_file_to_s3(
        file_bytes, prefix="kommersant/blanks", filename=filename, content_type=content_type,
    )
    return StoredFile.objects.create(
        bucket=bucket, key=key, filename=filename,
        content_type=content_type, size=len(file_bytes),
    )


def _attach(client, stored, employee):
    root = get_or_create_root(client)
    folder = _mk(client, root, "Публикации Коммерсантъ", "kommersant", 7)
    ClientFile.objects.create(
        folder=folder, stored_file=stored, name=stored.filename,
        size=stored.size or 0, content_type=stored.content_type, uploaded_by=employee,
    )


def generate_blank(publication, *, employee=None):
    """Собрать заявку: .docx (с факсимиле) + .pdf, подшить в папку клиента.

    Возвращает публикацию с заполненными blank_docx / blank_pdf.
    """
    if not (publication.text or "").strip():
        raise KommersantBlankError(
            "Сначала сформируйте текст сообщения — он печатается в бланке заявки."
        )

    ctx = build_blank_context(publication)
    docx_bytes = render_docx(_load_template(), ctx)

    proc = publication.procedure or _procedure_for(publication.case)
    am = proc.arbitr_manager if (proc and proc.arbitr_manager_id) else None
    docx_bytes = _apply_signature(docx_bytes, am)

    try:
        pdf_bytes = docx_to_pdf(docx_bytes)
    except Exception as exc:  # noqa: BLE001
        raise KommersantBlankError(f"Не удалось собрать PDF заявки: {exc}") from exc

    client = publication.case.service.client
    base = f"Заявка Коммерсантъ — {client.last_name or ''} {publication.type_label}".strip()[:120]
    docx_sf = _store(docx_bytes, filename=f"{base}.docx", content_type=DOCX_CT)
    pdf_sf = _store(pdf_bytes, filename=f"{base}.pdf", content_type="application/pdf")

    _attach(client, pdf_sf, employee)
    _attach(client, docx_sf, employee)

    publication.blank_docx = docx_sf
    publication.blank_pdf = pdf_sf
    publication.save(update_fields=["blank_docx", "blank_pdf", "updated_at"])

    try:
        client_log.invalidate_cache()
        client_log.record_action(
            client, "kommersant_blank_created",
            comment=f"Сформирована заявка на публикацию в «Коммерсантъ»: {publication.type_label}. "
                    f"Файл — в папке «Публикации Коммерсантъ».",
            employee=employee, stored_file=pdf_sf,
        )
    except Exception:
        log.exception("generate_blank: не удалось записать событийку")
    return publication


def save_text(publication, text: str, *, employee=None):
    """Сохранить текст сообщения (задача «сформировать текст»)."""
    text = (text or "").strip()
    if not text:
        raise KommersantBlankError("Текст сообщения пуст — сгенерируйте или введите его.")

    publication.text = text
    publication.title = publication.type_label
    publication.generated_at = timezone.now()
    if employee is not None and publication.created_by_id is None:
        publication.created_by = employee
    if publication.status == publication.STATUS_DRAFT:
        publication.status = publication.STATUS_GENERATED
    publication.save()

    try:
        client_log.invalidate_cache()
        client_log.record_action(
            publication.case.service.client, "kommersant_text_generated",
            comment=f"Сформирован текст сообщения для «Коммерсанта»: {publication.type_label}.",
            employee=employee,
        )
    except Exception:
        log.exception("save_text: не удалось записать событийку")
    return publication
